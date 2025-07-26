"""
Text Processing Module for WinForms Application
Handles file reading, text extraction, and word cloud generation
"""

import os
import io
import base64
from typing import Dict, List, Tuple, Any
from collections import Counter
import re

# Text processing libraries
import nltk
from wordcloud import WordCloud
import matplotlib.pyplot as plt
import matplotlib
matplotlib.use('Agg')  # Use non-interactive backend

# File processing libraries
import PyPDF2
from docx import Document
import markdown

# Download NLTK data if not present
def ensure_nltk_data():
    """Download required NLTK data if not present"""
    required_datasets = [
        ('tokenizers/punkt_tab', 'punkt_tab'),
        ('corpora/stopwords', 'stopwords')
    ]
    
    for data_path, dataset_name in required_datasets:
        try:
            nltk.data.find(data_path)
        except LookupError:
            print(f"Downloading NLTK dataset: {dataset_name}")
            nltk.download(dataset_name, quiet=True)

# Ensure NLTK data is available
ensure_nltk_data()

from nltk.corpus import stopwords
from nltk.tokenize import word_tokenize


def get_version() -> str:
    """
    Returns the version information for this module
    """
    return "Text Processor v1.0 - NLTK, WordCloud, PyPDF2, python-docx"


def extract_text_from_file(file_path: str) -> str:
    """
    Extracts text from various file formats
    
    Args:
        file_path: Path to the file to process
        
    Returns:
        Extracted text as string
        
    Raises:
        ValueError: If file type is not supported
        FileNotFoundError: If file doesn't exist
    """
    if not os.path.exists(file_path):
        raise FileNotFoundError(f"File not found: {file_path}")
    
    file_extension = os.path.splitext(file_path)[1].lower()
    
    try:
        if file_extension == '.txt':
            return _extract_from_txt(file_path)
        elif file_extension == '.pdf':
            return _extract_from_pdf(file_path)
        elif file_extension == '.docx':
            return _extract_from_docx(file_path)
        elif file_extension in ['.md', '.markdown']:
            return _extract_from_markdown(file_path)
        else:
            raise ValueError(f"Unsupported file type: {file_extension}")
    except Exception as e:
        raise ValueError(f"Error processing {file_extension} file: {str(e)}")


def _extract_from_txt(file_path: str) -> str:
    """Extract text from plain text file"""
    with open(file_path, 'r', encoding='utf-8', errors='ignore') as file:
        return file.read()


def _extract_from_pdf(file_path: str) -> str:
    """Extract text from PDF file using PyPDF2"""
    text = ""
    with open(file_path, 'rb') as file:
        pdf_reader = PyPDF2.PdfReader(file)
        
        for page_num in range(len(pdf_reader.pages)):
            page = pdf_reader.pages[page_num]
            text += page.extract_text() + "\n"
    
    return text


def _extract_from_docx(file_path: str) -> str:
    """Extract text from Word document"""
    doc = Document(file_path)
    text = ""
    
    # Extract text from paragraphs
    for paragraph in doc.paragraphs:
        text += paragraph.text + "\n"
    
    # Extract text from tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                text += cell.text + " "
        text += "\n"
    
    return text


def _extract_from_markdown(file_path: str) -> str:
    """Extract text from Markdown file by converting to plain text"""
    with open(file_path, 'r', encoding='utf-8', errors='ignore') as file:
        md_content = file.read()
    
    # Convert markdown to HTML, then strip HTML tags
    html = markdown.markdown(md_content)
    
    # Simple HTML tag removal (for basic cases)
    clean_text = re.sub(r'<[^>]+>', '', html)
    
    return clean_text


def clean_and_tokenize_text(text: str) -> List[str]:
    """
    Clean and tokenize text, removing stop words and non-alphabetic tokens
    
    Args:
        text: Raw text to process
        
    Returns:
        List of cleaned tokens
    """
    # Convert to lowercase and tokenize
    tokens = word_tokenize(text.lower())
    
    # Get English stop words
    stop_words = set(stopwords.words('english'))
    
    # Filter tokens: only alphabetic, not stop words, length > 2
    cleaned_tokens = [
        token for token in tokens 
        if token.isalpha() and token not in stop_words and len(token) > 2
    ]
    
    return cleaned_tokens


def generate_word_frequencies(tokens: List[str], top_n: int = 50) -> str:
    """
    Generate word frequency analysis
    
    Args:
        tokens: List of cleaned tokens
        top_n: Number of top words to return
        
    Returns:
        Formatted string with word frequencies
    """
    word_counts = Counter(tokens)
    
    # Get top N words
    top_words = word_counts.most_common(top_n)
    
    # Format as simple string that C# can easily parse
    lines = []
    for word, count in top_words:
        lines.append(f"{word}: {count}")
    
    return "\n".join(lines)


def create_word_cloud_image(tokens: List[str], width: int = 800, height: int = 400) -> str:
    """
    Create a word cloud image and return as base64 string
    
    Args:
        tokens: List of cleaned tokens
        width: Image width in pixels
        height: Image height in pixels
        
    Returns:
        Base64 encoded PNG image
    """
    # Join tokens back into text for WordCloud
    text_for_wordcloud = ' '.join(tokens)
    
    if not text_for_wordcloud.strip():
        # Create empty image if no text
        fig, ax = plt.subplots(figsize=(width/100, height/100))
        ax.text(0.5, 0.5, 'No text to display', 
               horizontalalignment='center', verticalalignment='center', 
               transform=ax.transAxes, fontsize=20)
        ax.set_xlim(0, 1)
        ax.set_ylim(0, 1)
        ax.axis('off')
    else:
        # Generate word cloud
        wordcloud = WordCloud(
            width=width,
            height=height,
            background_color='white',
            max_words=100,
            colormap='viridis',
            relative_scaling=0.5,
            random_state=42
        ).generate(text_for_wordcloud)
        
        # Create matplotlib figure
        fig, ax = plt.subplots(figsize=(width/100, height/100))
        ax.imshow(wordcloud, interpolation='bilinear')
        ax.axis('off')
    
    # Save to bytes buffer
    img_buffer = io.BytesIO()
    plt.savefig(img_buffer, format='png', bbox_inches='tight', 
                dpi=100, facecolor='white', edgecolor='none')
    img_buffer.seek(0)
    
    # Convert to base64
    img_base64 = base64.b64encode(img_buffer.getvalue()).decode('utf-8')
    
    # Clean up
    plt.close(fig)
    img_buffer.close()
    
    return img_base64


def process_file_complete(file_path: str) -> Dict[str, Any]:
    """
    Complete file processing pipeline: extract text, analyze, and create word cloud
    
    Args:
        file_path: Path to the file to process
        
    Returns:
        Dictionary containing analysis results
    """
    try:
        print(f"Processing file: {file_path}")
        
        # Ensure NLTK data is available before processing
        ensure_nltk_data()
        
        # Step 1: Extract text
        print("Step 1: Extracting text from file")
        raw_text = extract_text_from_file(file_path)
        
        if not raw_text.strip():
            return {
                "success": False,
                "error": "No text content found in file",
                "word_count": 0,
                "unique_words": 0,
                "frequencies": [],
                "word_cloud_base64": ""
            }
        
        print(f"Extracted {len(raw_text)} characters of text")
        
        # Step 2: Clean and tokenize
        print("Step 2: Cleaning and tokenizing text")
        tokens = clean_and_tokenize_text(raw_text)
        
        if not tokens:
            return {
                "success": False,
                "error": "No meaningful words found after processing",
                "word_count": 0,
                "unique_words": 0,
                "frequencies": [],
                "word_cloud_base64": ""
            }
        
        print(f"Processed {len(tokens)} tokens, {len(set(tokens))} unique words")
        
        # Step 3: Generate frequencies
        print("Step 3: Generating word frequencies")
        frequencies_text = generate_word_frequencies(tokens, top_n=50)
        
        # Step 4: Create word cloud
        print("Step 4: Creating word cloud image")
        word_cloud_base64 = create_word_cloud_image(tokens, width=800, height=400)
        
        # Step 5: Return results
        print("Step 5: Processing complete, returning results")
        return {
            "success": True,
            "error": "",
            "word_count": len(tokens),
            "unique_words": len(set(tokens)),
            "frequencies_text": frequencies_text,
            "word_cloud_base64": word_cloud_base64,
            "file_processed": os.path.basename(file_path)
        }
        
    except Exception as e:
        return {
            "success": False,
            "error": str(e),
            "word_count": 0,
            "unique_words": 0,
            "frequencies": [],
            "word_cloud_base64": ""
        }